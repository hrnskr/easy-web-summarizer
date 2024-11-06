import gradio as gr
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

from summarizer import load_document, setup_summarization_chain
from translator import setup_translator_chain
from yt_summarizer import check_link, summarize_video

#pdf zimbirtilari
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics


def summarize(url):
    if check_link(url):
        result = summarize_video(url)
    else:
        docs = load_document(url)
        llm_chain = setup_summarization_chain()
        result = llm_chain.run(docs)

    btn_translate = gr.Button("🇹🇷 Translate ", visible=True)
    btn_export_pdf = gr.Button("Export as PDF", visible=True)
    btn_export_docx = gr.Button("Export as DOCX", visible=True)

    return [result, btn_translate, btn_export_pdf, btn_export_docx]


def translate(text):
    llm_chain = setup_translator_chain()
    result = llm_chain.run(text)
    return result

# ~ def export_to_pdf2(text):
    # ~ pdf_file_path = "summary.pdf"
    # ~ c = canvas.Canvas(pdf_file_path, pagesize=letter)
    # ~ c.drawString(100, 750, text)  # Metni sayfada uygun bir yere yerleştir
    # ~ c.save()
    # ~ return pdf_file_path

def export_to_pdf(texts):
    doc_file_path = "summary.pdf"
    # PDF dosyasını oluştur
    pdf = SimpleDocTemplate(doc_file_path, pagesize=letter)
    elements = []

    # Türkçe karakter desteği için fontu 
    pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))  # Font dosyasının yolu
    styles = getSampleStyleSheet()
    
    # fontu ayarı
    normal_style = ParagraphStyle(
        name='Normal',
        fontName='DejaVuSans',
        fontSize=12,
        leading=15,  # Satır aralığı
    )

    # Metinleri birleştir
    # ~ print("=====")
    # ~ print(texts)
    # ~ print("=====")
    combined_text = "".join(texts)  # Metinleri HTML formatında birleştir
    # ~ print("vvvvvv")
    # ~ print(combined_text)
    # ~ print("^^^^^^")
    # Metni Paragraph ile oluştur ve PDF'ye ekle
    paragraph = Paragraph(combined_text, normal_style)
    elements.append(paragraph)

    # PDF dosyasını oluştur ve kaydet
    pdf.build(elements)
    return doc_file_path
    
def export_to_docx(text):
    doc = Document()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(text)
    doc_file_path = "summary.docx"
    doc.save(doc_file_path)
    return doc_file_path


with gr.Blocks() as demo:
    gr.Markdown(
        """# Cobanov Web and Video Summarizer
    Easily summarize any web page or YouTube video with a single click."""
    )

    with gr.Row():
        with gr.Column():
            url = gr.Text(label="URL", placeholder="Enter URL here")

            btn_generate = gr.Button("Generate")

            summary = gr.Markdown(label="Summary")
            btn_translate = gr.Button(visible=False)
            btn_export_pdf = gr.Button("Export as PDF", visible=False)
            btn_export_docx = gr.Button("Export as DOCX", visible=False)

    gr.Examples(
        [
            "https://cobanov.dev/haftalik-bulten/hafta-13",
            "https://bawolf.substack.com/p/embeddings-are-a-good-starting-point",
            "https://www.youtube.com/watch?v=4pOpQwiUVXc",
        ],
        inputs=[url],
    )
    gr.Markdown(
        """
        ```
        Model: llama3-8b
        Author: Mert Cobanov
        Contact: mertcobanov@gmail.com
        Repo: github.com/mertcobanov/easy-web-summarizer
        ```"""
    )
    btn_generate.click(summarize, inputs=[url], outputs=[summary, btn_translate, btn_export_pdf, btn_export_docx])
    btn_translate.click(translate, inputs=[summary], outputs=[summary])
    # ~ btn_export_pdf.click(export_to_pdf, inputs=[summary], outputs=None)
    # ~ btn_export_docx.click(export_to_docx, inputs=[summary], outputs=None)
    btn_export_pdf.click(export_to_pdf, inputs=[summary], outputs=gr.File(label="Download PDF"))  # PDF için dosya döndür
    btn_export_docx.click(export_to_docx, inputs=[summary], outputs=gr.File(label="Download DOCX"))  # DOCX için dosya döndür

demo.launch(server_name="0.0.0.0")
